"""
Document Loader Module for GeM TenderLens.
Extracts structured text and page metadata from PDF, DOCX, XLSX, CSV, and Email files.
Detects scanned/image-only PDFs and marks them as manual review required.
"""

import os
import re
import email
from email import policy
from typing import Dict, List, Any, Tuple
from utils_logger import get_logger

logger = get_logger(__name__)

class DocumentLoader:
    """Multi-format document loader with text-accessibility validation."""

    @staticmethod
    def clean_english_text(text: str) -> str:
        """Strips Devanagari script, corrupt Hindi font characters, and orphan slashes for pure English output."""
        if not text:
            return ""
        # Remove Devanagari / Indic script Unicode ranges (\u0900-\u0DFF)
        text = re.sub(r'[\u0900-\u0DFF]+', '', text)
        # Remove non-ASCII characters resulting from Hindi font decoding artifacts
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)
        # Remove junk backticks, tildes, stray symbols
        text = re.sub(r'[`~]+', '', text)
        # Clean up orphan slashes (e.g. ' / ' or '/ ')
        text = re.sub(r'\s*/\s*', ' ', text)
        # Normalize whitespace
        lines = [re.sub(r'\s+', ' ', line).strip() for line in text.splitlines()]
        return '\n'.join([line for line in lines if line])

    @staticmethod
    def load_document(file_path: str) -> Dict[str, Any]:
        """
        Loads document content based on file extension.
        
        Returns dict with:
            - filename (str)
            - ext (str)
            - pages (List[Dict[str, Any]]): page_number, content, is_scanned
            - is_scanned (bool): overall scan flag
            - error (Optional[str])
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return {
                "filename": os.path.basename(file_path),
                "ext": "",
                "pages": [],
                "is_scanned": False,
                "error": f"File not found: {file_path}"
            }

        filename = os.path.basename(file_path)
        ext = os.path.splitext(filename)[1].lower()
        logger.info(f"Loading document: {filename} (extension: {ext})")

        loaders = {
            ".pdf": DocumentLoader._load_pdf,
            ".docx": DocumentLoader._load_docx,
            ".doc": DocumentLoader._load_docx,
            ".xlsx": DocumentLoader._load_excel,
            ".xls": DocumentLoader._load_excel,
            ".csv": DocumentLoader._load_csv,
            ".eml": DocumentLoader._load_email_or_txt,
            ".msg": DocumentLoader._load_email_or_txt,
            ".txt": DocumentLoader._load_email_or_txt,
        }

        try:
            loader_fn = loaders.get(ext)
            if loader_fn:
                result = loader_fn(file_path, filename)
                # Centralized text cleaning for all document formats
                for page in result.get("pages", []):
                    if page.get("content"):
                        page["content"] = DocumentLoader.clean_english_text(page["content"])
                return result
            
            logger.warning(f"Unsupported extension {ext} for file {filename}")
            return {
                "filename": filename,
                "ext": ext,
                "pages": [{"page_number": 1, "content": f"Unsupported file format: {ext}", "is_scanned": False}],
                "is_scanned": False,
                "error": f"Unsupported file format: {ext}"
            }
        except Exception as e:
            logger.exception(f"Error loading document {file_path}: {e}")
            return {
                "filename": filename,
                "ext": ext,
                "pages": [],
                "is_scanned": False,
                "error": str(e)
            }

    @staticmethod
    def _load_pdf(file_path: str, filename: str) -> Dict[str, Any]:
        """Extracts text from PDF using PyMuPDF (fitz) and flags scanned pages."""
        pages = []
        overall_scanned = False
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            total_pages = len(doc)
            logger.info(f"Opened PDF '{filename}' with {total_pages} pages")

            for page_idx in range(total_pages):
                page = doc[page_idx]
                raw_text = page.get_text("text").strip()
                text = DocumentLoader.clean_english_text(raw_text)
                
                # Check if page is image-only/scanned (very few characters extracted)
                is_scanned = len(raw_text) < 30 or len(text) < 15
                if is_scanned:
                    overall_scanned = True
                    text += "\n[MANUAL REVIEW REQUIRED: Page appears to be scanned or image-only document]"

                pages.append({
                    "page_number": page_idx + 1,
                    "content": text,
                    "is_scanned": is_scanned
                })
            doc.close()
        except Exception as e:
            logger.exception(f"Failed to parse PDF {file_path}: {e}")
            raise

        return {
            "filename": filename,
            "ext": ".pdf",
            "pages": pages,
            "is_scanned": overall_scanned,
            "error": None
        }

    @staticmethod
    def _load_docx(file_path: str, filename: str) -> Dict[str, Any]:
        """Extracts text from DOCX files using python-docx."""
        pages = []
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for p in doc.paragraphs:
                if p.text.strip():
                    full_text.append(p.text.strip())
            
            # Read tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        full_text.append(f"Table row: {row_text}")

            content = "\n".join(full_text)
            pages.append({"page_number": 1, "content": content, "is_scanned": False})
        except Exception as e:
            logger.exception(f"Failed to parse DOCX {file_path}: {e}")
            raise

        return {
            "filename": filename,
            "ext": ".docx",
            "pages": pages,
            "is_scanned": False,
            "error": None
        }

    @staticmethod
    def _load_excel(file_path: str, filename: str) -> Dict[str, Any]:
        """Extracts text from Excel files (.xlsx / .xls) """
        pages = []
        ext = os.path.splitext(filename)[1].lower()
        try:
            import pandas as pd
            xls = pd.ExcelFile(file_path)
            page_num = 1
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                sheet_content = f"--- Sheet: {sheet_name} ---\n" + df.to_string(index=False)
                pages.append({"page_number": page_num, "content": sheet_content, "is_scanned": False})
                page_num += 1
        except Exception as e:
            logger.exception(f"Failed to parse Excel file {file_path}: {e}")
            return {
                "filename": filename,
                "ext": ext,
                "pages": [],
                "is_scanned": False,
                "error": str(e)
            }

        return {
            "filename": filename,
            "ext": ext,
            "pages": pages,
            "is_scanned": False,
            "error": None
        }

    @staticmethod
    def _load_csv(file_path: str, filename: str) -> Dict[str, Any]:
        """Extracts text from CSV files."""
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            content = df.to_string(index=False)
            return {
                "filename": filename,
                "ext": ".csv",
                "pages": [{"page_number": 1, "content": content, "is_scanned": False}],
                "is_scanned": False,
                "error": None
            }
        except Exception as e:
            logger.exception(f"Failed to parse CSV file {file_path}: {e}")
            raise

    @staticmethod
    def _load_email_or_txt(file_path: str, filename: str) -> Dict[str, Any]:
        """Extracts text from raw email files (.eml) or text files (.txt)."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                raw_content = f.read()

            if file_path.endswith(".eml"):
                msg = email.message_from_string(raw_content, policy=policy.default)
                subject = msg.get("subject", "No Subject")
                sender = msg.get("from", "Unknown Sender")
                date_hdr = msg.get("date", "")
                
                body = ""
                if msg.is_multipart():
                    for part in msg.walk():
                        if part.get_content_type() == "text/plain":
                            body += part.get_content()
                else:
                    body = msg.get_content()
                
                content = f"From: {sender}\nDate: {date_hdr}\nSubject: {subject}\n\n{body}"
            else:
                content = raw_content

            return {
                "filename": filename,
                "ext": os.path.splitext(filename)[1].lower(),
                "pages": [{"page_number": 1, "content": content, "is_scanned": False}],
                "is_scanned": False,
                "error": None
            }
        except Exception as e:
            logger.exception(f"Failed to parse Email/TXT file {file_path}: {e}")
            raise
